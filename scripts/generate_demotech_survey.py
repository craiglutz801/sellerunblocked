#!/usr/bin/env python3
"""Generate DemoTech SellerUnblocked survey Word doc and 500-response CSV."""

import csv
import random
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / ".tmp_pydeps"))
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents" / "demotech-survey"
OUT_DIR.mkdir(parents=True, exist_ok=True)

random.seed(42)

# ---------------------------------------------------------------------------
# Survey definition
# ---------------------------------------------------------------------------

COMPANY = "DemoTech"
COMPANY_TAGLINE = "Software that makes sales demos unbelievable and amazing."

SECTIONS = [
    {
        "title": "Section A — About You",
        "intro": "These questions help us segment results. Individual responses remain confidential; leadership sees aggregated patterns only.",
        "questions": [
            {"id": "A1", "type": "text", "text": "Employee ID (assigned by HR — do not use your name)."},
            {"id": "A2", "type": "single", "text": "Job title", "options": [
                "SDR", "BDR", "Commercial AE", "Mid-Market AE", "Enterprise AE", "Strategic AE",
                "Sales Manager", "Regional Sales Director", "VP Sales", "CRO",
                "Sales Engineer", "Solutions Consultant", "Sales Enablement Manager",
                "RevOps Analyst", "RevOps Manager", "GTM Marketing Manager",
            ]},
            {"id": "A3", "type": "single", "text": "Role level", "options": [
                "Individual Contributor", "Front-line Manager", "Director+", "Executive",
            ]},
            {"id": "A4", "type": "single", "text": "Team / segment", "options": [
                "Commercial", "Mid-Market", "Enterprise", "Strategic", "Corporate / GTM",
            ]},
            {"id": "A5", "type": "single", "text": "Primary work location", "options": [
                "San Francisco, CA (HQ)", "New York, NY", "Austin, TX", "Denver, CO",
                "London, UK", "Toronto, Canada", "Remote — US West", "Remote — US East",
                "Remote — EMEA", "Remote — Canada",
            ]},
            {"id": "A6", "type": "single", "text": "Tenure at DemoTech", "options": [
                "0–6 months", "6–12 months", "1–2 years", "2–5 years", "5+ years",
            ]},
        ],
    },
    {
        "title": "Section B — Sales Process Effectiveness",
        "intro": "How well the sales process helps you move deals forward.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "B1", "type": "likert5", "text": "I clearly understand our sales stages and what is required to advance a deal."},
            {"id": "B2", "type": "likert5", "text": "Our qualification criteria help me focus on winnable opportunities."},
            {"id": "B3", "type": "likert5", "text": "Handoffs between SDR/BDR and Account Executives are smooth and do not create deal friction."},
            {"id": "B4", "type": "likert5", "text": "Our demo-to-close process is well defined for my segment."},
            {"id": "B5", "type": "likert5", "text": "Forecasting expectations and stage definitions are consistent across the revenue org."},
        ],
    },
    {
        "title": "Section C — Technology & Tool Experience",
        "intro": "Whether tools and systems support selling rather than slow you down.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "C1", "type": "likert5", "text": "Our CRM is configured to support how I sell, not slow me down."},
            {"id": "C2", "type": "likert5", "text": "DemoTech's demo platform is easy for me to use in live selling situations."},
            {"id": "C3", "type": "likert5", "text": "I can find the right demo assets, templates, and content quickly when preparing for a customer meeting."},
            {"id": "C4", "type": "likert5", "text": "Our sales tech stack (Outreach, Gong, Slack, etc.) works well together for my daily workflow."},
            {"id": "C5", "type": "likert5", "text": "Reporting and pipeline views give me accurate information without manual workarounds."},
        ],
    },
    {
        "title": "Section D — Sales Enablement Quality",
        "intro": "Training, content, and resources that help you perform.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "D1", "type": "likert5", "text": "Onboarding prepared me to run effective demos within my first 90 days."},
            {"id": "D2", "type": "likert5", "text": "Enablement content is relevant to the deals I am actually working."},
            {"id": "D3", "type": "likert5", "text": "I get enough coaching and practice on objection handling and competitive positioning."},
            {"id": "D4", "type": "likert5", "text": "Product updates are communicated in time for me to adjust my selling approach."},
            {"id": "D5", "type": "likert5", "text": "Playbooks and talk tracks reflect what actually works with our buyers."},
        ],
    },
    {
        "title": "Section E — Marketing Alignment",
        "intro": "How well marketing supports pipeline and seller needs.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "E1", "type": "likert5", "text": "Marketing messaging aligns with what I hear from buyers in the field."},
            {"id": "E2", "type": "likert5", "text": "MQLs and inbound leads match our ICP and are worth my time."},
            {"id": "E3", "type": "likert5", "text": "Case studies, ROI content, and proof points are available for my buyer personas."},
            {"id": "E4", "type": "likert5", "text": "Campaign themes and product launches are coordinated with field feedback."},
        ],
    },
    {
        "title": "Section F — Lead Quality & Pipeline Support",
        "intro": "Whether leads and pipeline are set up for success.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "F1", "type": "likert5", "text": "I have enough qualified pipeline to hit my quota."},
            {"id": "F2", "type": "likert5", "text": "Territory and account coverage design sets me up for success."},
            {"id": "F3", "type": "likert5", "text": "Outbound and inbound support from SDR/BDR teams is effective."},
            {"id": "F4", "type": "likert5", "text": "Partner-sourced and channel opportunities are well qualified before they reach me."},
        ],
    },
    {
        "title": "Section G — Compensation Clarity",
        "intro": "How clear and fair compensation and incentives are.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "G1", "type": "likert5", "text": "I understand how my compensation plan works and what drives my earnings."},
            {"id": "G2", "type": "likert5", "text": "Compensation changes are communicated clearly and with enough lead time."},
            {"id": "G3", "type": "likert5", "text": "I feel the comp plan rewards the behaviors that grow the business."},
            {"id": "G4", "type": "likert5", "text": "SPIFFs and accelerators are transparent and easy to track."},
        ],
    },
    {
        "title": "Section H — Internal Support Systems",
        "intro": "Deal support, operations, and cross-functional help.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "H1", "type": "likert5", "text": "Deal desk, legal, and security review turnaround is predictable."},
            {"id": "H2", "type": "likert5", "text": "I can get timely help from Sales Engineers when I need technical support on a deal."},
            {"id": "H3", "type": "likert5", "text": "RevOps provides reliable answers when I have process or data questions."},
            {"id": "H4", "type": "likert5", "text": "Customer Success and Support teams help me protect and expand accounts."},
        ],
    },
    {
        "title": "Section I — Organizational Alignment",
        "intro": "Strategy, goals, and priorities understood across the org.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "I1", "type": "likert5", "text": "Company strategy and GTM priorities are clearly communicated."},
            {"id": "I2", "type": "likert5", "text": "Product, Sales, and Marketing are aligned on how we win in the market."},
            {"id": "I3", "type": "likert5", "text": "Cross-functional teams understand the pressure and pace of selling."},
            {"id": "I4", "type": "likert5", "text": "Leadership decisions reflect an understanding of front-line selling realities."},
        ],
    },
    {
        "title": "Section J — Overall Seller Experience",
        "intro": "Your overall confidence and satisfaction in the revenue organization.",
        "scale": "1 = Strongly disagree · 2 = Disagree · 3 = Neutral · 4 = Agree · 5 = Strongly agree",
        "questions": [
            {"id": "J1", "type": "likert5", "text": "I am satisfied working in DemoTech's revenue organization."},
            {"id": "J2", "type": "likert5", "text": "I would recommend DemoTech as a place to sell to other sales professionals."},
            {"id": "J3", "type": "likert5", "text": "I have confidence I can hit my quota this year."},
            {"id": "J4", "type": "likert5", "text": "My manager provides consistent coaching that helps me win deals."},
        ],
    },
    {
        "title": "Section K — Sales Stage Confidence",
        "intro": "How confident you feel executing each stage of the sales cycle.",
        "scale": "1 = Not at all confident · 2 = Slightly confident · 3 = Moderately confident · 4 = Confident · 5 = Very confident",
        "questions": [
            {"id": "K1", "type": "likert5", "text": "Discover — finding and engaging the right buyers."},
            {"id": "K2", "type": "likert5", "text": "Qualify — determining fit, urgency, and decision process."},
            {"id": "K3", "type": "likert5", "text": "Demo — delivering a compelling, tailored product demonstration."},
            {"id": "K4", "type": "likert5", "text": "Propose / Negotiate — pricing, ROI, and commercial terms."},
            {"id": "K5", "type": "likert5", "text": "Close — getting contracts signed and deals across the line."},
        ],
    },
    {
        "title": "Section L — Open Feedback",
        "intro": "Optional but valuable. Responses are anonymized before leadership review.",
        "questions": [
            {"id": "L1", "type": "open", "text": "What is the single biggest thing slowing you down from selling more effectively?"},
            {"id": "L2", "type": "open", "text": "What one change would most improve your ability to win deals?"},
            {"id": "L3", "type": "open", "text": "Anything else you want leadership to understand about your experience? (Optional)"},
        ],
    },
]

# Flatten likert question ids for CSV generation
LIKERT_IDS = [
    q["id"]
    for sec in SECTIONS
    for q in sec["questions"]
    if q.get("type") == "likert5"
]

OPEN_IDS = ["L1", "L2", "L3"]

# ---------------------------------------------------------------------------
# Word document
# ---------------------------------------------------------------------------

def build_word_doc() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"{COMPANY} Revenue Organization Assessment", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(COMPANY_TAGLINE)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Survey instrument · SellerUnblocked · Cycle Q2 {date.today().year}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Purpose: This confidential assessment measures the operating environment for DemoTech's "
        "revenue team — process, tools, enablement, alignment, and support — so leadership can "
        "prioritize fixes that help sellers win. Estimated completion time: 12–15 minutes."
    )

    doc.add_heading("Instructions", level=1)
    for bullet in [
        "Answer honestly. There are no right or wrong answers.",
        "Your individual responses are confidential. Leadership sees aggregated trends, not attributable quotes.",
        "Sections B–K use a 1–5 scale unless noted otherwise.",
        "Open-text responses may be paraphrased or clustered before sharing with leadership.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    q_num = 0
    for section in SECTIONS:
        doc.add_page_break()
        doc.add_heading(section["title"], level=1)
        if section.get("intro"):
            doc.add_paragraph(section["intro"])
        if section.get("scale"):
            p = doc.add_paragraph()
            p.add_run("Response scale: ").bold = True
            p.add_run(section["scale"])

        for q in section["questions"]:
            q_num += 1
            doc.add_paragraph()
            pq = doc.add_paragraph()
            pq.add_run(f"{q['id']}. ").bold = True
            pq.add_run(q["text"])

            if q["type"] == "likert5":
                doc.add_paragraph("  ○ 1   ○ 2   ○ 3   ○ 4   ○ 5", style="List Bullet")
            elif q["type"] == "single":
                for opt in q["options"]:
                    doc.add_paragraph(f"  ○ {opt}", style="List Bullet")
            elif q["type"] == "text":
                doc.add_paragraph("_" * 60)
            elif q["type"] == "open":
                doc.add_paragraph("(Open text — 2–4 sentences encouraged)")
                for _ in range(3):
                    doc.add_paragraph("_" * 60)

    doc.add_page_break()
    doc.add_heading("Thank You", level=1)
    doc.add_paragraph(
        "Your feedback helps DemoTech build a sales organization that matches the quality of "
        "the demos we deliver to customers. A personalized enablement summary may be shared with "
        "you privately in a follow-up — leadership will not see individual attribution."
    )

    out_path = OUT_DIR / "DemoTech-Revenue-Organization-Survey.docx"
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Synthetic response generation
# ---------------------------------------------------------------------------

TITLE_WEIGHTS = [
    ("SDR", 0.12),
    ("BDR", 0.08),
    ("Commercial AE", 0.18),
    ("Mid-Market AE", 0.16),
    ("Enterprise AE", 0.14),
    ("Strategic AE", 0.04),
    ("Sales Manager", 0.09),
    ("Regional Sales Director", 0.03),
    ("VP Sales", 0.015),
    ("CRO", 0.005),
    ("Sales Engineer", 0.05),
    ("Solutions Consultant", 0.03),
    ("Sales Enablement Manager", 0.02),
    ("RevOps Analyst", 0.015),
    ("RevOps Manager", 0.01),
    ("GTM Marketing Manager", 0.01),
]

TITLE_TO_LEVEL = {
    "SDR": "Individual Contributor",
    "BDR": "Individual Contributor",
    "Commercial AE": "Individual Contributor",
    "Mid-Market AE": "Individual Contributor",
    "Enterprise AE": "Individual Contributor",
    "Strategic AE": "Individual Contributor",
    "Sales Engineer": "Individual Contributor",
    "Solutions Consultant": "Individual Contributor",
    "Sales Manager": "Front-line Manager",
    "Regional Sales Director": "Director+",
    "VP Sales": "Executive",
    "CRO": "Executive",
    "Sales Enablement Manager": "Individual Contributor",
    "RevOps Analyst": "Individual Contributor",
    "RevOps Manager": "Front-line Manager",
    "GTM Marketing Manager": "Individual Contributor",
}

TITLE_TO_SEGMENT = {
    "SDR": "Commercial",
    "BDR": "Mid-Market",
    "Commercial AE": "Commercial",
    "Mid-Market AE": "Mid-Market",
    "Enterprise AE": "Enterprise",
    "Strategic AE": "Strategic",
    "Sales Manager": "Commercial",
    "Regional Sales Director": "Enterprise",
    "VP Sales": "Corporate / GTM",
    "CRO": "Corporate / GTM",
    "Sales Engineer": "Enterprise",
    "Solutions Consultant": "Mid-Market",
    "Sales Enablement Manager": "Corporate / GTM",
    "RevOps Analyst": "Corporate / GTM",
    "RevOps Manager": "Corporate / GTM",
    "GTM Marketing Manager": "Corporate / GTM",
}

LOCATIONS = [
    "San Francisco, CA (HQ)", "New York, NY", "Austin, TX", "Denver, CO",
    "London, UK", "Toronto, Canada", "Remote — US West", "Remote — US East",
    "Remote — EMEA", "Remote — Canada",
]

TENURE_BANDS = ["0–6 months", "6–12 months", "1–2 years", "2–5 years", "5+ years"]
TENURE_MONTHS = {
    "0–6 months": (1, 6),
    "6–12 months": (7, 12),
    "1–2 years": (13, 24),
    "2–5 years": (25, 60),
    "5+ years": (61, 120),
}

# Dimension base scores (org-level tendencies for DemoTech)
DIM_BASE = {
    "B": 3.6,  # process — decent but handoff friction
    "C": 3.9,  # tools — they sell demo software, generally good
    "D": 3.5,  # enablement — mixed
    "E": 3.2,  # marketing alignment — weaker
    "F": 3.4,  # pipeline
    "G": 3.7,  # comp
    "H": 3.1,  # support — deal desk pain
    "I": 3.5,  # alignment
    "J": 3.6,  # overall
    "K": 3.7,  # stage confidence — demo stage high
}

OPEN_TEXT_POOL = {
    "process": [
        "Stage definitions change every quarter and my manager interprets them differently than RevOps.",
        "SDR-to-AE handoffs lose context — I re-discover pain points on the first call every time.",
        "Forecasting feels like a political exercise instead of an honest pipeline review.",
        "We have six different definitions of 'qualified' depending on who you ask.",
    ],
    "tools": [
        "CRM fields are a graveyard — half the required data doesn't help me sell.",
        "Ironically our internal demo environment crashes more than I'd like on live calls.",
        "Finding the right demo template in Seismic takes longer than building one from scratch.",
        "Gong insights are great but nothing feeds back into Salesforce automatically.",
    ],
    "enablement": [
        "Onboarding was strong on product but weak on enterprise procurement cycles.",
        "Competitive battlecards are three months stale the moment a competitor ships something.",
        "I want more live role-play on CFO objections, not another recorded webinar.",
        "Enablement content is built for Commercial but I'm Enterprise — I adapt everything.",
    ],
    "marketing": [
        "Marketing promises 'demo transformation in 30 days' and buyers expect magic on call one.",
        "MQLs from the last campaign were mostly students and consultants, not buyers.",
        "We need ROI calculators for mid-market CFOs — I build them in Google Sheets weekly.",
        "Product marketing launches features before sales knows how to position them.",
    ],
    "pipeline": [
        "Territory carve-up left me with a lot of long-dormant accounts and no warm inbound.",
        "Inbound volume dropped after we changed the website — outbound isn't picking up the slack.",
        "Partner referrals arrive with no context and unrealistic close timelines.",
        "My patch is geographically huge for an Enterprise rep — travel eats selling time.",
    ],
    "comp": [
        "SPIFF rules changed mid-quarter and nobody could explain the clawback logic.",
        "Multi-year deal splits are unclear when CS owns the expansion motion.",
        "Accelerators kick in too late — I don't feel rewarded for landing lighthouse logos.",
        "Comp plan deck was 40 slides; I still don't know how renewals affect my number.",
    ],
    "support": [
        "Legal review on enterprise deals takes 2–3 weeks and champions go dark.",
        "SE coverage is thin — I wait three days for a technical win on competitive deals.",
        "Deal desk gives different pricing guidance depending on who is on shift.",
        "Security questionnaire responses are copy-paste from last year and buyers catch it.",
    ],
    "alignment": [
        "Product roadmap slides don't match what I hear in discovery calls.",
        "We say we're PLG-plus-sales but the field motion still feels like pure enterprise.",
        "All-hands celebrates product velocity; sellers are drowning in half-baked releases.",
        "Remote reps feel last to know when GTM strategy shifts — HQ decides, field executes.",
    ],
    "positive": [
        "Our demo platform is genuinely best-in-class — buyers feel it on the first call.",
        "My manager runs tight pipeline reviews with real coaching, not just inspection.",
        "The Commercial team culture is strong — people share what works on Slack.",
        "RevOps fixed forecasting dashboards this quarter and it saved me hours every week.",
    ],
}

CHANGE_SUGGESTIONS = [
    "Standardize SDR-to-AE handoff notes with required fields in Salesforce.",
    "Staff one more SE per Enterprise pod for competitive bake-offs.",
    "Publish a single source of truth for stage definitions and enforce it in Gong.",
    "Refresh ROI and CFO content for mid-market — stop making AEs rebuild decks.",
    "Cap legal review SLA at 5 business days for deals under $150K ARR.",
    "Run monthly live objection labs instead of async enablement modules.",
    "Tighten MQL criteria — fewer leads, higher quality.",
    "Give remote reps the same GTM briefings HQ gets in hallway conversations.",
    "Simplify the comp plan to one page with worked examples.",
    "Build demo environment stability into our own dogfooding SLA.",
]


def weighted_choice(pairs):
    items, weights = zip(*pairs)
    return random.choices(items, weights=weights, k=1)[0]


def clamp_score(val):
    return max(1, min(5, round(val)))


def likert_score(base, noise=0.0, bias=0):
    return clamp_score(random.gauss(base + bias + noise, 0.85))


def generate_respondent(idx: int) -> dict:
    title = weighted_choice(TITLE_WEIGHTS)
    level = TITLE_TO_LEVEL[title]
    segment = TITLE_TO_SEGMENT[title]
    # Managers/directors may cover different segments
    if title == "Sales Manager":
        segment = random.choice(["Commercial", "Mid-Market", "Enterprise"])
    elif title == "Regional Sales Director":
        segment = random.choice(["Enterprise", "Strategic", "Mid-Market"])

    location = random.choice(LOCATIONS)
    tenure = random.choices(
        TENURE_BANDS,
        weights=[0.14, 0.16, 0.22, 0.28, 0.20],
        k=1,
    )[0]
    lo, hi = TENURE_MONTHS[tenure]
    tenure_mo = random.randint(lo, hi)

    row = {
        "response_id": f"DT-{date.today().year}Q2-{idx:04d}",
        "company": COMPANY,
        "employee_id": f"DT-{10000 + idx}",
        "title": title,
        "role_level": level,
        "team_segment": segment,
        "location": location,
        "tenure_band": tenure,
        "tenure_months": tenure_mo,
        "response_date": f"2025-{random.randint(4, 5):02d}-{random.randint(1, 28):02d}",
    }

    # Persona biases
    tenure_bias = {
        "0–6 months": -0.35,
        "6–12 months": -0.15,
        "1–2 years": 0.0,
        "2–5 years": 0.1,
        "5+ years": 0.05,
    }[tenure]

    remote_bias = -0.2 if location.startswith("Remote") else 0.05
    ent_support_bias = -0.45 if segment == "Enterprise" else 0.0
    sdr_pipeline_bias = -0.3 if title in ("SDR", "BDR") else 0.0

    for qid in LIKERT_IDS:
        dim = qid[0]
        base = DIM_BASE[dim]
        bias = tenure_bias + remote_bias

        if dim == "H":
            bias += ent_support_bias
        if dim == "F":
            bias += sdr_pipeline_bias
        if dim == "C" and qid == "C2":
            bias += 0.35  # demo platform generally liked
        if dim == "K" and qid == "K3":
            bias += 0.5  # demo stage confidence high at DemoTech
        if dim == "K" and qid == "K4":
            bias += ent_support_bias * 0.5  # negotiate harder for enterprise
        if title in ("Sales Enablement Manager", "RevOps Analyst", "RevOps Manager"):
            bias += 0.25  # GTM ops slightly more positive on process/tools
        if title in ("VP Sales", "CRO"):
            bias += 0.4

        row[qid] = likert_score(base, bias=bias)

    # Open text — pick themes based on lowest dimension scores
    dim_avgs = {}
    for dim in "BCDEFGHIJK":
        ids = [q for q in LIKERT_IDS if q.startswith(dim)]
        dim_avgs[dim] = sum(row[i] for i in ids) / len(ids)

    worst_dims = sorted(dim_avgs, key=dim_avgs.get)[:2]
    theme_map = {
        "B": "process", "C": "tools", "D": "enablement", "E": "marketing",
        "F": "pipeline", "G": "comp", "H": "support", "I": "alignment",
    }
    themes = [theme_map.get(d, "alignment") for d in worst_dims if d in theme_map]
    if random.random() < 0.15:
        themes = ["positive"]

    row["L1"] = random.choice(OPEN_TEXT_POOL[themes[0]])
    row["L2"] = random.choice(CHANGE_SUGGESTIONS)
    if random.random() < 0.65:
        row["L3"] = random.choice(OPEN_TEXT_POOL[themes[-1] if len(themes) > 1 else themes[0]])
    else:
        row["L3"] = ""

    return row


def build_csv(n: int = 500) -> Path:
    demo_fields = [
        "response_id", "company", "employee_id", "title", "role_level",
        "team_segment", "location", "tenure_band", "tenure_months", "response_date",
    ]
    fieldnames = demo_fields + LIKERT_IDS + OPEN_IDS

    rows = [generate_respondent(i + 1) for i in range(n)]
    out_path = OUT_DIR / "DemoTech-Survey-Responses-500.csv"
    with out_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    return out_path


def main():
    doc_path = build_word_doc()
    csv_path = build_csv(500)
    print(f"Word doc: {doc_path}")
    print(f"CSV:      {csv_path}")
    print(f"Rows:     500")
    print(f"Likert Q: {len(LIKERT_IDS)}")


if __name__ == "__main__":
    main()
